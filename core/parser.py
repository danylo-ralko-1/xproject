"""Multi-format file parser for requirements ingestion.

Supported formats:
  - PDF (.pdf)           → pdfplumber
  - Word (.docx)         → python-docx
  - Excel (.xlsx, .xls)  → openpyxl
  - CSV (.csv)           → csv module
  - Email (.eml)         → email module
  - Plain text (.txt, .md, .rtf) → direct read
  - Images (.png, .jpg, .jpeg, .gif, .webp) → base64 for Claude vision
"""

import csv
import email
import base64
import mimetypes
from pathlib import Path
from dataclasses import dataclass, field


# Extensions grouped by parser type
TEXT_EXTS = {".txt", ".md", ".rtf", ".text"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
EXCEL_EXTS = {".xlsx", ".xls"}
CSV_EXTS = {".csv"}
EMAIL_EXTS = {".eml"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}

ALL_SUPPORTED = TEXT_EXTS | PDF_EXTS | DOCX_EXTS | EXCEL_EXTS | CSV_EXTS | EMAIL_EXTS | IMAGE_EXTS

# Context window threshold: ~150K tokens at 4 chars/token.
# Leaves ~50K tokens for system prompts, conversation history, and working memory.
CONTEXT_CHAR_THRESHOLD = 600_000


@dataclass
class ParsedFile:
    """Result of parsing a single file."""
    filename: str
    format: str          # "text", "pdf", "docx", "excel", "csv", "email", "image"
    text: str = ""       # Extracted text content
    is_image: bool = False
    image_base64: str = ""       # Base64-encoded image data (for vision)
    image_media_type: str = ""   # MIME type of image
    metadata: dict = field(default_factory=dict)  # Extra info (sheet names, email headers, etc.)
    error: str = ""      # Error message if parsing failed


def parse_file(filepath: Path) -> ParsedFile:
    """
    Parse a single file and extract its content.
    Routes to the appropriate parser based on extension.
    """
    ext = filepath.suffix.lower()

    if ext in TEXT_EXTS:
        return _parse_text(filepath)
    elif ext in PDF_EXTS:
        return _parse_pdf(filepath)
    elif ext in DOCX_EXTS:
        return _parse_docx(filepath)
    elif ext in EXCEL_EXTS:
        return _parse_excel(filepath)
    elif ext in CSV_EXTS:
        return _parse_csv(filepath)
    elif ext in EMAIL_EXTS:
        return _parse_email(filepath)
    elif ext in IMAGE_EXTS:
        return _parse_image(filepath)
    else:
        return ParsedFile(
            filename=filepath.name,
            format="unknown",
            error=f"Unsupported format: {ext}",
        )


def parse_directory(directory: Path) -> list[ParsedFile]:
    """
    Parse all supported files in a directory (recursively).
    Returns list of ParsedFile objects, sorted by filename.
    """
    results = []
    if not directory.exists():
        return results

    for f in sorted(directory.rglob("*")):
        if not f.is_file():
            continue
        if f.name.startswith("."):
            continue
        if f.suffix.lower() not in ALL_SUPPORTED:
            results.append(ParsedFile(
                filename=f.name,
                format="unknown",
                error=f"Skipped unsupported format: {f.suffix}",
            ))
            continue
        results.append(parse_file(f))

    return results


def build_context(parsed_files: list[ParsedFile]) -> tuple[str, list[dict]]:
    """
    Build a combined context string from all parsed files.

    Returns:
        (text_context, image_blocks)
        - text_context: combined text from all non-image files
        - image_blocks: list of dicts for Claude vision API
          [{"type": "image", "source": {"type": "base64", ...}}]
    """
    text_parts = []
    image_blocks = []

    for pf in parsed_files:
        if pf.error:
            continue

        if pf.is_image:
            image_blocks.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": pf.image_media_type,
                    "data": pf.image_base64,
                },
            })
            # Also note the image in text context for reference
            text_parts.append(f"--- [{pf.filename}] (image attached for visual review) ---")
        elif pf.text.strip():
            header = f"--- [{pf.filename}] ({pf.format}) ---"
            text_parts.append(f"{header}\n{pf.text.strip()}")

    text_context = "\n\n".join(text_parts)
    return text_context, image_blocks


# --- Individual parsers ---

def _parse_text(filepath: Path) -> ParsedFile:
    """Parse plain text files."""
    try:
        text = filepath.read_text(encoding="utf-8", errors="replace")
        return ParsedFile(filename=filepath.name, format="text", text=text)
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="text", error=str(e))


def _parse_pdf(filepath: Path) -> ParsedFile:
    """Parse PDF files using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return ParsedFile(
            filename=filepath.name, format="pdf",
            error="pdfplumber not installed. Run: pip install pdfplumber",
        )

    try:
        pages = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                parts = [text]
                for table in tables:
                    parts.append(_table_to_text(table))
                pages.append(f"[Page {i+1}]\n{chr(10).join(parts)}")

        return ParsedFile(
            filename=filepath.name,
            format="pdf",
            text="\n\n".join(pages),
            metadata={"page_count": len(pages)},
        )
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="pdf", error=str(e))


def _parse_docx(filepath: Path) -> ParsedFile:
    """Parse Word documents using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return ParsedFile(
            filename=filepath.name, format="docx",
            error="python-docx not installed. Run: pip install python-docx",
        )

    try:
        doc = Document(filepath)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Preserve heading structure
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                prefix = "#" * int(level) if level.isdigit() else "##"
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                parts.append(_rows_to_text(rows))

        return ParsedFile(
            filename=filepath.name,
            format="docx",
            text="\n\n".join(parts),
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
        )
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="docx", error=str(e))


def _parse_excel(filepath: Path) -> ParsedFile:
    """Parse Excel files using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ParsedFile(
            filename=filepath.name, format="excel",
            error="openpyxl not installed. Run: pip install openpyxl",
        )

    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # Skip completely empty rows
                if all(cell is None for cell in row):
                    continue
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                rows.append(cells)

            if rows:
                parts.append(f"[Sheet: {sheet_name}]")
                parts.append(_rows_to_text(rows))

        wb.close()
        return ParsedFile(
            filename=filepath.name,
            format="excel",
            text="\n\n".join(parts),
            metadata={"sheets": wb.sheetnames},
        )
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="excel", error=str(e))


def _parse_csv(filepath: Path) -> ParsedFile:
    """Parse CSV files."""
    try:
        rows = []
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            # Sniff delimiter
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
                reader = csv.reader(f, dialect)
            except csv.Error:
                reader = csv.reader(f)

            for row in reader:
                cells = [cell.strip() for cell in row]
                if any(cells):
                    rows.append(cells)

        return ParsedFile(
            filename=filepath.name,
            format="csv",
            text=_rows_to_text(rows),
            metadata={"row_count": len(rows)},
        )
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="csv", error=str(e))


def _parse_email(filepath: Path) -> ParsedFile:
    """Parse .eml email files."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            msg = email.message_from_file(f)

        headers = {
            "from": msg.get("From", ""),
            "to": msg.get("To", ""),
            "date": msg.get("Date", ""),
            "subject": msg.get("Subject", ""),
        }

        # Extract body
        body_parts = []
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_parts.append(payload.decode("utf-8", errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body_parts.append(payload.decode("utf-8", errors="replace"))

        header_text = (
            f"From: {headers['from']}\n"
            f"To: {headers['to']}\n"
            f"Date: {headers['date']}\n"
            f"Subject: {headers['subject']}"
        )
        body_text = "\n\n".join(body_parts)
        full_text = f"{header_text}\n\n{body_text}"

        return ParsedFile(
            filename=filepath.name,
            format="email",
            text=full_text,
            metadata=headers,
        )
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="email", error=str(e))


def _parse_image(filepath: Path) -> ParsedFile:
    """Parse image files → base64 for Claude vision."""
    try:
        raw = filepath.read_bytes()
        b64 = base64.b64encode(raw).decode("ascii")

        mime = mimetypes.guess_type(filepath.name)[0] or "image/png"
        # Normalize common types
        mime_map = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".webp": "image/webp",
        }
        media_type = mime_map.get(filepath.suffix.lower(), mime)

        return ParsedFile(
            filename=filepath.name,
            format="image",
            is_image=True,
            image_base64=b64,
            image_media_type=media_type,
            metadata={"size_bytes": len(raw)},
        )
    except Exception as e:
        return ParsedFile(filename=filepath.name, format="image", error=str(e))


# --- Context management helpers ---

def estimate_tokens(text: str) -> int:
    """Rough token estimate: ~4 characters per token."""
    return len(text) // 4


def build_sections(parsed_files: list[ParsedFile]) -> list[dict]:
    """
    Build per-file section metadata for large document sets.

    When the combined context exceeds CONTEXT_CHAR_THRESHOLD, each file gets its
    own numbered markdown section that can be read independently by skills.

    Returns a list of dicts:
        [{"index": 1, "filename": "001_RFP-Document.pdf.md",
          "source": "RFP-Document.pdf", "format": "pdf",
          "chars": 42300, "estimated_tokens": 10575, "content": "..."}]
    """
    sections = []
    idx = 0
    for pf in parsed_files:
        if pf.error or pf.is_image or not pf.text.strip():
            continue
        idx += 1
        safe_name = pf.filename.replace(" ", "-")
        section_filename = f"{idx:03d}_{safe_name}.md"
        content = f"# {pf.filename} ({pf.format})\n\n{pf.text.strip()}"
        sections.append({
            "index": idx,
            "filename": section_filename,
            "source": pf.filename,
            "format": pf.format,
            "chars": len(content),
            "estimated_tokens": estimate_tokens(content),
            "content": content,
        })
    return sections


# --- Table formatting helpers ---

def _table_to_text(table: list) -> str:
    """Convert a pdfplumber table (list of lists) to readable text."""
    if not table:
        return ""
    rows = []
    for row in table:
        cells = [str(cell).strip() if cell else "" for cell in row]
        rows.append(cells)
    return _rows_to_text(rows)


def _rows_to_text(rows: list[list[str]]) -> str:
    """Convert rows of cells into an aligned text table."""
    if not rows:
        return ""

    # Calculate column widths
    max_cols = max(len(r) for r in rows)
    # Pad rows to same length
    padded = [r + [""] * (max_cols - len(r)) for r in rows]

    col_widths = []
    for c in range(max_cols):
        w = max(len(padded[r][c]) for r in range(len(padded)))
        col_widths.append(min(w, 60))  # Cap at 60 chars

    lines = []
    for i, row in enumerate(padded):
        cells = [cell[:60].ljust(col_widths[j]) for j, cell in enumerate(row)]
        lines.append(" | ".join(cells))
        if i == 0:
            lines.append("-+-".join("-" * w for w in col_widths))

    return "\n".join(lines)
